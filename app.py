import streamlit as st
import docx
from docx.shared import Pt, Inches
import pypdf
import re
import io

def split_mcq_options(line):
    # Fix any option letter glued to the previous word and inject a space
    cleaned_line = re.sub(r'([^\s])([A-D])\.\s', r'\1 \2. ', line)
    
    matches = list(re.finditer(r'([A-D])\.\s', cleaned_line))
    if not matches: return []
    
    options = []
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i+1].start() if i + 1 < len(matches) else len(cleaned_line)
        options.append(cleaned_line[start:end].strip())
    return options

def add_options_grid(doc, options):
    if not options: return
    max_len = max(len(o) for o in options)
    num_cols = 4 if (len(options) == 4 and max_len < 20) else 2
    num_rows = (len(options) + num_cols - 1) // num_cols
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = False
    col_width = Inches(6.2 / num_cols)
    for row in table.rows:
        for cell in row.cells: cell.width = col_width
            
    for index, option in enumerate(options):
        row_idx = index // num_cols
        col_idx = index % num_cols
        cell = table.rows[row_idx].cells[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(option)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = docx.oxml.OxmlElement('w:tcBorders')
            for b_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                b = docx.oxml.OxmlElement(f'w:{b_name}')
                b.set(docx.oxml.ns.qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)

def replace_tags_in_tables(tables, subject, grade, date, marks, exam_title):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "[SUB]" in paragraph.text:
                        # SURGICAL TEXT SWAP: Keeps existing runs/formatting/borders completely safe
                        for run in paragraph.runs:
                            if "[SUB]" in run.text:
                                run.text = run.text.replace("[SUB]", subject)
                    elif "[GRD]" in paragraph.text:
                        for run in paragraph.runs:
                            if "[GRD]" in run.text:
                                run.text = run.text.replace("[GRD]", grade)
                    elif "[DT]" in paragraph.text:
                        for run in paragraph.runs:
                            if "[DT]" in run.text:
                                run.text = run.text.replace("[DT]", date)
                    elif "[MRK]" in paragraph.text:
                        for run in paragraph.runs:
                            if "[MRK]" in run.text:
                                run.text = run.text.replace("[MRK]", marks)
                    elif "[EXM]" in paragraph.text:
                        # Clear paragraph runs cleanly for the multi-line Exam Box title input
                        paragraph.text = ""
                        run = paragraph.add_run(exam_title)
                        run.font.name = 'Times New Roman'
                        run.font.bold = True

def process_paper(pdf_file, template_path, subject, grade, date, marks, exam_title):
    reader = pypdf.PdfReader(pdf_file)
    pdf_lines = []
    for page in reader.pages:
        text = page.extract_text()
        for line in text.split('\n'):
            if line.strip(): pdf_lines.append(line.strip())

    doc = docx.Document(template_path)
    
    # Process text styling swaps inside the main template grid body
    replace_tags_in_tables(doc.tables, subject, grade, date, marks, exam_title)
    
    # Process text styling swaps inside page headers
    for section in doc.sections:
        replace_tags_in_tables(section.header.tables, subject, grade, date, marks, exam_title)
        if section.first_page_header:
            replace_tags_in_tables(section.first_page_header.tables, subject, grade, date, marks, exam_title)

    # Clear instructions text below the header box
    while len(doc.paragraphs) > 0:
        p_to_remove = doc.paragraphs[-1]
        p_to_remove._element.getparent().remove(p_to_remove._element)

    current_options = []
    for line in pdf_lines:
        clean_line = line.strip()
        upper_line = clean_line.upper()
        
        if "SUBJECT:" in upper_line or "STANDARD:" in upper_line or "QUESTION PAPER" in upper_line or "MARKS:" in upper_line:
            continue
        if re.match(r'^\(\d+\)$', clean_line):
            continue

        has_explicit_options = bool(re.search(r'[A-D]\.\s', clean_line))
        is_broken_continuation = len(current_options) > 0 and not has_explicit_options and not re.match(r'^\d+\.', clean_line) and "SECTION -" not in upper_line and "►" not in clean_line

        if is_broken_continuation:
            current_options[-1] = current_options[-1] + " " + clean_line
            continue

        if not has_explicit_options and current_options:
            add_options_grid(doc, current_options)
            current_options = []

        if "SECTION -" in upper_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
        elif "►" in clean_line or "ANSWER THE FOLLOWING" in upper_line or "CHOOSE THE CORRECT" in upper_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            mark_match = re.search(r'\(\d{2}\)$', clean_line)
            if mark_match:
                mark_text = mark_match.group(0)
                instruction_text = clean_line[:mark_match.start()].strip()
                run = p.add_run(instruction_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                run_mark = p.add_run(f"\t{mark_text}")
                run_mark.font.name = 'Times New Roman'
                run_mark.font.size = Pt(14)
                run_mark.bold = True
            else:
                run = p.add_run(clean_line)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
        elif has_explicit_options:
            extracted = split_mcq_options(clean_line)
            current_options.extend(extracted)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    if current_options:
        add_options_grid(doc, current_options)

    # Footer sign-off text entry
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(12)
    
    run_footer = p_footer.add_run("ALL THE BEST")
    run_footer.font.name = 'Aptos Display'
    run_footer.font.size = Pt(20)
    run_footer.bold = True
    run_footer.italic = True

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# --- STREAMLIT USER INTERFACE ---
st.set_page_config(page_title="Exam Paper Formatter", page_icon="📄")
st.title("📄 K N Patel Exam Converter")
st.write("Fill out the exam variables below, upload the raw PDF, and generate your custom formatted Word file instantly.")

col1, col2 = st.columns(2)
with col1:
    subject_input = st.text_input("Subject Name", value="S.P.C.C.")
    grade_input = st.text_input("Grade / Standard", value="11th GSEB")
    exam_title_input = st.text_input("Examination Title Box", value="FA – 1 Examination\n[2026-27]")
with col2:
    date_input = st.text_input("Date of Exam", value="24-06-2026")
    marks_input = st.text_input("Total Marks", value="30")

uploaded_pdf = st.file_uploader("Upload Raw PDF Question Paper", type=["pdf"])

if uploaded_pdf:
    if st.button("✨ Convert Document Now"):
        with st.spinner("Processing custom formatting rules..."):
            try:
                output_docx = process_paper(
                    uploaded_pdf, 
                    "template.docx", 
                    subject_input, 
                    grade_input, 
                    date_input, 
                    marks_input,
                    exam_title_input
                )
                st.success("🎉 Conversion Complete!")
                st.download_button(
                    label="📥 Download Formatted Word Document",
                    data=output_docx,
                    file_name=f"Formatted_{subject_input}_Paper.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Something went wrong: {e}")
